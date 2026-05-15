"""
Unit tests for document_tools module
测试文档处理工具的各项功能

"""

import asyncio
import json
import os
import tempfile
import unittest
from pathlib import Path

# 尝试导入依赖库，如果没有则跳过测试
try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


class TestDocumentTools(unittest.TestCase):
    """文档处理工具测试类"""

    # ==================== 工具运行适配器 ====================
    @staticmethod
    def _run_tool(func, *args, **kwargs):
        """
        适配工具函数同步/异步两种实现：
        - 如果被 @FunctionHub.tool 装饰，函数会被替换为 async，需要 await
        - 如果是普通函数，直接返回结果
        """
        try:
            result = func(*args, **kwargs)
            if asyncio.iscoroutine(result):
                return asyncio.run(result)
            return result
        except TypeError:
            # 某些实现可能返回协程函数对象
            return asyncio.run(func(*args, **kwargs))

    @classmethod
    def setUpClass(cls):
        """创建测试数据目录"""
        cls.test_dir = tempfile.mkdtemp(prefix="doc_tools_test_")
        cls.test_files = {}

        # 创建测试PDF
        if PYMUPDF_AVAILABLE:
            cls._create_test_pdf()

        # 创建测试Word文档
        if DOCX_AVAILABLE:
            cls._create_test_docx()

        # 创建测试Excel文件
        if OPENPYXL_AVAILABLE:
            cls._create_test_excel()

    @classmethod
    def _create_test_pdf(cls):
        """创建测试用的PDF文件"""
        pdf_path = Path(cls.test_dir) / "test_document.pdf"
        doc = fitz.open()

        # 第1页 - 文本内容
        page1 = doc.new_page(width=595, height=842)  # A4尺寸
        page1.insert_text((50, 50), "Test Document", fontsize=20)
        page1.insert_text((50, 100), "This is page 1 with some test content.")
        page1.insert_text((50, 130), "测试中文内容：这是第一页。")

        # 第2页 - 更多文本
        page2 = doc.new_page()
        page2.insert_text((50, 50), "Page 2", fontsize=20)
        page2.insert_text((50, 100), "More content on page 2.")

        # 第3页 - 简单内容
        page3 = doc.new_page()
        page3.insert_text((50, 50), "Page 3", fontsize=20)

        doc.save(pdf_path)
        doc.close()

        cls.test_files["pdf"] = str(pdf_path)

        # 创建第二个PDF用于合并测试
        pdf_path2 = Path(cls.test_dir) / "test_document2.pdf"
        doc2 = fitz.open()
        page = doc2.new_page()
        page.insert_text((50, 50), "Second PDF for merge test")
        doc2.save(pdf_path2)
        doc2.close()

        cls.test_files["pdf2"] = str(pdf_path2)

    @classmethod
    def _create_test_docx(cls):
        """创建测试用的Word文档"""
        docx_path = Path(cls.test_dir) / "test_document.docx"
        doc = Document()

        # 添加标题
        doc.add_heading("Test Document", 0)

        # 添加段落
        doc.add_paragraph("This is the first paragraph.")
        doc.add_paragraph("这是第二段，包含中文内容。")
        doc.add_paragraph("Third paragraph with more content.")

        # 添加表格
        table = doc.add_table(rows=3, cols=3)
        table.style = "Light Grid Accent 1"

        # 填充表格
        headers = ["Name", "Age", "City"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        table.rows[1].cells[0].text = "Alice"
        table.rows[1].cells[1].text = "30"
        table.rows[1].cells[2].text = "Beijing"

        table.rows[2].cells[0].text = "Bob"
        table.rows[2].cells[1].text = "25"
        table.rows[2].cells[2].text = "Shanghai"

        doc.save(docx_path)
        cls.test_files["docx"] = str(docx_path)

    @classmethod
    def _create_test_excel(cls):
        """创建测试用的Excel文件"""
        excel_path = Path(cls.test_dir) / "test_document.xlsx"
        wb = Workbook()

        # 第一个工作表
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["Name", "Age", "City"])
        ws1.append(["Alice", 30, "Beijing"])
        ws1.append(["Bob", 25, "Shanghai"])
        ws1.append(["Charlie", 35, "Guangzhou"])

        # 第二个工作表
        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["Product", "Price", "Quantity"])
        ws2.append(["Apple", 5.0, 100])
        ws2.append(["Banana", 3.0, 150])

        wb.save(excel_path)
        cls.test_files["excel"] = str(excel_path)

    @classmethod
    def tearDownClass(cls):
        """清理测试文件"""
        import shutil

        if os.path.exists(cls.test_dir):
            shutil.rmtree(cls.test_dir)

    # ==================== PDF 测试 ====================

    @unittest.skipIf(not PYMUPDF_AVAILABLE, "PyMuPDF not installed")
    def test_extract_pdf_text(self):
        """测试PDF文本提取"""
        from function_hubs.document_tools import extract_pdf_text

        result_str = self._run_tool(extract_pdf_text, self.test_files["pdf"])
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(result["total_pages"], 3)
        self.assertTrue(len(result["pages"]) > 0)
        self.assertIn("Test Document", result["pages"][0]["text"])

    @unittest.skipIf(not PYMUPDF_AVAILABLE, "PyMuPDF not installed")
    def test_extract_pdf_text_with_range(self):
        """测试指定页码范围提取PDF文本"""
        from function_hubs.document_tools import extract_pdf_text

        # 测试单页
        result_str = self._run_tool(
            extract_pdf_text, self.test_files["pdf"], page_range="1"
        )
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(result["extracted_pages"], 1)

        # 测试范围
        result_str = self._run_tool(
            extract_pdf_text, self.test_files["pdf"], page_range="1-2"
        )
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(result["extracted_pages"], 2)

    @unittest.skipIf(not PYMUPDF_AVAILABLE, "PyMuPDF not installed")
    def test_get_pdf_info(self):
        """测试获取PDF信息"""
        from function_hubs.document_tools import get_pdf_info

        result_str = self._run_tool(get_pdf_info, self.test_files["pdf"])
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(result["document_properties"]["page_count"], 3)
        self.assertFalse(result["document_properties"]["is_encrypted"])
        self.assertIn("file_info", result)
        self.assertIn("content_statistics", result)

    @unittest.skipIf(not PYMUPDF_AVAILABLE, "PyMuPDF not installed")
    def test_merge_pdfs(self):
        """测试PDF合并"""
        from function_hubs.document_tools import merge_pdfs

        output_path = os.path.join(self.test_dir, "merged.pdf")
        pdf_list = [self.test_files["pdf"], self.test_files["pdf2"]]

        result_str = self._run_tool(merge_pdfs, pdf_list, output_path)
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertTrue(os.path.exists(output_path))
        self.assertEqual(result["total_pages"], 4)  # 3 + 1

    @unittest.skipIf(not PYMUPDF_AVAILABLE, "PyMuPDF not installed")
    def test_split_pdf(self):
        """测试PDF拆分"""
        from function_hubs.document_tools import split_pdf

        output_dir = os.path.join(self.test_dir, "split_output")
        split_ranges = ["1", "2-3"]

        result_str = self._run_tool(
            split_pdf,
            self.test_files["pdf"],
            split_ranges,
            output_dir,
        )
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(len(result["files"]), 2)

        # 验证文件存在
        for file_info in result["files"]:
            self.assertTrue(os.path.exists(file_info["path"]))

    @unittest.skipIf(not PYMUPDF_AVAILABLE, "PyMuPDF not installed")
    def test_extract_pdf_images(self):
        """测试PDF图像提取（空PDF无图像）"""
        from function_hubs.document_tools import extract_pdf_images

        output_dir = os.path.join(self.test_dir, "images_output")

        result_str = self._run_tool(
            extract_pdf_images, self.test_files["pdf"], output_dir
        )
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        # 测试PDF没有图像，所以应该为0
        self.assertEqual(result["image_count"], 0)

    # ==================== Word 测试 ====================

    @unittest.skipIf(not DOCX_AVAILABLE, "python-docx not installed")
    def test_read_docx(self):
        """测试读取Word文档"""
        from function_hubs.document_tools import read_docx

        result_str = self._run_tool(read_docx, self.test_files["docx"])
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertGreater(result["statistics"]["paragraph_count"], 0)
        self.assertGreater(result["statistics"]["table_count"], 0)

        # 验证段落内容
        paragraphs = result["paragraphs"]
        self.assertTrue(any("first paragraph" in p["text"] for p in paragraphs))

    @unittest.skipIf(not DOCX_AVAILABLE, "python-docx not installed")
    def test_extract_docx_text(self):
        """测试提取Word文档纯文本"""
        from function_hubs.document_tools import extract_docx_text

        result_str = self._run_tool(extract_docx_text, self.test_files["docx"])
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertIn("first paragraph", result["text"])
        self.assertGreater(result["length"], 0)

    # ==================== Excel 测试 ====================

    @unittest.skipIf(not OPENPYXL_AVAILABLE, "openpyxl not installed")
    def test_read_excel(self):
        """测试读取Excel文件"""
        from function_hubs.document_tools import read_excel

        result_str = self._run_tool(read_excel, self.test_files["excel"])
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(result["statistics"]["row_count"], 3)
        self.assertEqual(result["headers"], ["Name", "Age", "City"])

        # 验证数据
        self.assertEqual(result["rows"][0][0], "Alice")

    @unittest.skipIf(not OPENPYXL_AVAILABLE, "openpyxl not installed")
    def test_read_excel_specific_sheet(self):
        """测试读取Excel特定工作表"""
        from function_hubs.document_tools import read_excel

        result_str = self._run_tool(
            read_excel, self.test_files["excel"], sheet_name="Sheet2"
        )
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(result["sheet_name"], "Sheet2")
        self.assertEqual(result["headers"], ["Product", "Price", "Quantity"])

    @unittest.skipIf(not OPENPYXL_AVAILABLE, "openpyxl not installed")
    def test_list_excel_sheets(self):
        """测试列出Excel工作表"""
        from function_hubs.document_tools import list_excel_sheets

        result_str = self._run_tool(list_excel_sheets, self.test_files["excel"])
        result = json.loads(result_str)

        self.assertTrue(result["success"])
        self.assertEqual(result["sheet_count"], 2)

        sheet_names = [s["name"] for s in result["sheets"]]
        self.assertIn("Sheet1", sheet_names)
        self.assertIn("Sheet2", sheet_names)

    # ==================== 通用功能测试 ====================

    def test_detect_document_format(self):
        """测试文档格式检测"""
        from function_hubs.document_tools import detect_document_format

        if "pdf" in self.test_files:
            result_str = self._run_tool(detect_document_format, self.test_files["pdf"])
            result = json.loads(result_str)

            self.assertTrue(result["success"])
            self.assertEqual(result["format_info"]["type"], "PDF")
            self.assertTrue(result["format_info"]["supported"])

    def test_parse_page_range(self):
        """测试页码范围解析"""
        from function_hubs.document_tools import _parse_page_range

        # 测试所有页
        pages = _parse_page_range(None, 10)
        self.assertEqual(pages, list(range(10)))

        # 测试单页
        pages = _parse_page_range("1", 10)
        self.assertEqual(pages, [0])

        # 测试范围
        pages = _parse_page_range("1-3", 10)
        self.assertEqual(pages, [0, 1, 2])

        # 测试逗号分隔
        pages = _parse_page_range("1,3,5", 10)
        self.assertEqual(pages, [0, 2, 4])

        # 测试组合
        pages = _parse_page_range("1-3,5,7-9", 10)
        self.assertEqual(pages, [0, 1, 2, 4, 6, 7, 8])

    # ==================== 错误处理测试 ====================

    def test_file_not_exist_error(self):
        """测试文件不存在错误处理"""
        from function_hubs.document_tools import extract_pdf_text, read_docx, read_excel

        non_exist_file = "/tmp/non_exist_file.pdf"

        # PDF
        if PYMUPDF_AVAILABLE:
            result = json.loads(self._run_tool(extract_pdf_text, non_exist_file))
            self.assertIn("error", result)

        # Word
        if DOCX_AVAILABLE:
            result = json.loads(self._run_tool(read_docx, non_exist_file))
            self.assertIn("error", result)

        # Excel
        if OPENPYXL_AVAILABLE:
            result = json.loads(self._run_tool(read_excel, non_exist_file))
            self.assertIn("error", result)

    @unittest.skipIf(not PYMUPDF_AVAILABLE, "PyMuPDF not installed")
    def test_invalid_page_range(self):
        """测试无效页码范围"""
        from function_hubs.document_tools import extract_pdf_text

        # 超出范围的页码
        result_str = self._run_tool(
            extract_pdf_text, self.test_files["pdf"], page_range="100-200"
        )
        result = json.loads(result_str)

        # 应该返回错误或空结果
        self.assertTrue("error" in result or result["extracted_pages"] == 0)


class TestDocumentToolsIntegration(unittest.TestCase):
    """集成测试：测试工具在OxyGent框架中的集成"""

    def test_tool_registration(self):
        """测试工具是否正确注册到FunctionHub"""
        try:
            from function_hubs.document_tools import document_tools

            # 验证是FunctionHub实例
            from oxygent.oxy import FunctionHub

            self.assertIsInstance(document_tools, FunctionHub)

            # 验证工具名称
            self.assertEqual(document_tools.name, "document_tools")

            # 验证至少有一些工具被注册
            # FunctionHub 使用 func_dict 保存注册函数
            self.assertGreater(len(document_tools.func_dict), 0)

        except ImportError as e:
            self.skipTest(f"Cannot import document_tools: {e}")

    def test_tool_import_from_preset_tools(self):
        """测试从preset_tools模块导入"""
        try:
            from function_hubs import document_tools

            # 验证document_tools在function_hubs中
            self.assertTrue(hasattr(document_tools, "document_tools"))

        except ImportError as e:
            self.skipTest(f"Cannot import function_hubs: {e}")


def run_tests():
    """运行所有测试"""
    # 创建测试套件
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    # 添加测试
    suite.addTests(loader.loadTestsFromTestCase(TestDocumentTools))
    suite.addTests(loader.loadTestsFromTestCase(TestDocumentToolsIntegration))

    # 运行测试
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    return result


if __name__ == "__main__":
    # 检查依赖
    print("检查依赖库...")
    print(f"PyMuPDF: {'✓ 已安装' if PYMUPDF_AVAILABLE else '✗ 未安装'}")
    print(f"pdfplumber: {'✓ 已安装' if PDFPLUMBER_AVAILABLE else '✗ 未安装'}")
    print(f"python-docx: {'✓ 已安装' if DOCX_AVAILABLE else '✗ 未安装'}")
    print(f"openpyxl: {'✓ 已安装' if OPENPYXL_AVAILABLE else '✗ 未安装'}")
    print("\n开始测试...\n")

    # 运行测试
    result = run_tests()

    # 打印总结
    print("\n" + "=" * 70)
    print("测试总结:")
    print(f"  运行: {result.testsRun}")
    print(f"  成功: {result.testsRun - len(result.failures) - len(result.errors)}")
    print(f"  失败: {len(result.failures)}")
    print(f"  错误: {len(result.errors)}")
    print(f"  跳过: {len(result.skipped)}")
    print("=" * 70)
